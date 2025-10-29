import os
from pathlib import Path
import re
from datetime import datetime
from typing import List, Dict, Any, Union
import logging
import PyPDF2
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:

    SUPPORTED_FORMATS = {'pdf', 'doc', 'docx', 'txt'}
    
    def __init__(
        self,
        chunk_size: int,
        overlap_size: int,
        embedding_model: str
    ):
        self._chunk_size = chunk_size
        self._overlap_size = overlap_size
        self._embedding_model = embedding_model
        self._model = None
    
    def _load_embedding_model(self):
        if self._model is None:
            logger.info(f"Loading embedding model: {self._embedding_model}")
            self._model = SentenceTransformer(self._embedding_model)
            logger.info("Embedding model loaded successfully.")
        return self._model

    def extract_text(self, document: Union[str, Path, bytes]) -> str:
        if isinstance(document, bytes):
            file_type = file_type.lower()
            if file_type == 'pdf':
                reader = PyPDF2.PdfReader(document)
                text = ''.join(page.extract_text() for page in reader.pages)
            elif file_type in {'doc', 'docx'}:
                doc = DocxDocument(document)
                text = '\n'.join(para.text for para in doc.paragraphs)
            elif file_type == 'txt':
                for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        text = document.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # Fallback with error handling
                    text = document.decode('utf-8', errors='replace')
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        else:
            file_path = Path(document)
            file_type = file_path.suffix[1:].lower()
            if file_type == 'pdf':
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ''.join(page.extract_text() for page in reader.pages)
            elif file_type in {'doc', 'docx'}:
                doc = DocxDocument(file_path)
                text = '\n'.join(para.text for para in doc.paragraphs)
            elif file_type == 'txt':
                # Try multiple encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            text = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # Fallback with error handling
                    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                        text = f.read()
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        return text
    
    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'\s{2,}', ' ', text)
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()
    
    def _find_sentence_boundary(self, text: str) -> int:

        sentence_endings = ['. ','! ','? ','.\n', '!\n', '?\n']

        last_position = -1

        for ending in sentence_endings:
            pos = text.rfind(ending)
            if pos > last_position:
                last_position = pos + len(ending)
        
        return last_position
        
    def chunk_text(self, text):
        text = self._clean_text(text)

        chunks = []
        start = 0
        chunk_id = 0

        while start < len(text):
            end = start + self._chunk_size
            if end > len(text):
                search_start = max(end-100, start)
                sentence_end = self._find_sentence_boundary(text[search_start:end])
                if sentence_end != -1:
                    end = search_start + sentence_end

            chunk = text[start:end].strip()
            if chunk:
                chunks.append({
                'chunk_id': chunk_id,
                'text': chunk,
                'start_char': start,
                'end_char': end,
                'chunk_size': len(chunk)
                })
                chunk_id += 1

            start = end - self._overlap_size

            if start <= (chunks[-1]['start_char'] if chunks else 0):
                start = end
        
        logger.info(f"created {len(chunks)} chunks from {len(text)} characters.")

        return chunks

    def extract_metadata(self, document: Union[str, Path]) -> Dict[str, Any]:
        file_path = Path(document)

        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        stat_info = file_path.stat()

        metadata = {
            'source': file_path.name,
            'filetype': file_path.suffix.lower(),
            'date_ingested': datetime.now().isoformat()
        }

        return metadata

    def embed_text(self, text_chunks: List[Union[str, Dict[str, Any]]]) -> List[List[float]]:
        if not text_chunks:
            logger.warning("No chunks provided for embedding")
            return []
        else:
            texts = text_chunks
        
        if isinstance(text_chunks[0], dict):
            texts = [chunk['text'] for chunk in text_chunks]
        else:
            texts = text_chunks

        model = self._load_embedding_model()
        logger.info(f"Generating embeddings for {len(texts)} chunks.")

        embeddings = model.encode(
            texts,
            show_progress_bar=True,
            convert_to_numpy=True
        )

        embeddings_list = embeddings.tolist()

        logger.info(f"Generated embeddings with shape: {len(embeddings_list)} x {len(embeddings_list[0])}")
        
        return embeddings_list
        

    def process_document(
            self,
            document: Union[str, Path],
        ) -> Dict[str, Any]:
        # Full processing pipeline
        logger.info(f"Processing document: {document}")

        try:
            metadata = self.extract_metadata(document)

            text = self.extract_text(document)
            if not text.strip():
                raise ValueError("Extracted text is empty.")

            chunks = self.chunk_text(text)
            if not chunks:
                raise ValueError("No chunks were created from the document text.")

            
            embeddings = self.embed_text(chunks)

            metadata['num_chunks'] = len(chunks)

            result = {
                'metadata': metadata,
                'text': text,
                'chunks': chunks,
                'embeddings': embeddings
            }

            logger.info(f"Document processing complete: {document}: {len(chunks)} chunks created.")

            return result
        
        except Exception as e:
            logger.error(f"Error processing document {document}: {e}")
            raise