from pathlib import Path
import re
from typing import List, Union
import PyPDF2
from docx import Document as DocxDocument
from app.config import settings


def extract_text(document: Union[str, Path, bytes]) -> str:
    '''
    Extract text from a document

    Params:
        document - the document to extract text from.

    Returns:
        the cleaned extracted text.
    '''
    if isinstance(document, bytes):
        file_type = file_type.lower()
        if file_type == 'pdf':
            reader = PyPDF2.PdfReader(document)
            text = ''.join(page.extract_text() for page in reader.pages)
        elif file_type in {'doc', 'docx'}:
            doc = DocxDocument(document)
            text = '\n'.join(para.text for para in doc.paragraphs)
        elif file_type == 'txt':
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = document.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # Fallback with error handling
                text = document.decode('utf-8', errors='replace')
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    else:
        file_path = Path(document)
        file_type = file_path.suffix[1:].lower()
        if file_type == 'pdf':
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ''.join(page.extract_text() for page in reader.pages)
        elif file_type in {'doc', 'docx'}:
            doc = DocxDocument(file_path)
            text = '\n'.join(para.text for para in doc.paragraphs)
        elif file_type == 'txt':
            # Try multiple encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # Fallback with error handling
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    return text

def clean_text(text: str) -> str:
    '''
    Clean some input text.

    Params:
        text - the text to clean

    Returns:
        The cleaned text.
    '''
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'\s{2,}', ' ', text)
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    return text.strip()

def find_sentence_boundary(text: str) -> int:
    '''
    Find the boundary between sentences in a piece of text.

    Params:
        text - the text in which to find the sentence boundary
    
    Returns:
        the index of the sentence boundary in the text string.
    '''
    sentence_endings = ['. ','! ','? ','.\n', '!\n', '?\n']
    last_position = -1
    for ending in sentence_endings:
        pos = text.rfind(ending)
        if pos > last_position:
            last_position = pos + len(ending)
    
    return last_position
    
def chunk_text(text: str)-> List[str]:
    '''
    Chunk a text into a series of smaller pieces of text with overlap.

    Params:
        text - The text to chunk.

    Returns:
        A List of chunk texts
    '''
    text = clean_text(text)
    chunks = []
    start = 0
    while start < len(text):
        end = start + settings.CHUNK_SIZE
        if end > len(text):
            search_start = max(end-100, start)
            sentence_end = find_sentence_boundary(text[search_start:end])
            if sentence_end != -1:
                end = search_start + sentence_end
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - settings.CHUNK_OVERLAP
        if start <= (chunks[-1]['start_char'] if chunks else 0):
            start = end
    return chunks